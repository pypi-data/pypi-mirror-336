# %% #@ Libraries imports

import pandas as pd
import numpy as np
import requests
import json
from io import BytesIO
import copy
import zipfile as zp
import xlsxwriter
import xml.etree.ElementTree as ET
import uuid
from types import FunctionType
from .classes import Form
import docx as dcx
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from re import findall
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.shared import Pt

# %% #@ Functions

def save_to_excel(data, filename="output.xlsx", column_width=25, include_index=False, row_colours={0: "#D8E4BC", 1: "#C5D9F1"}, row_bold=[0], row_wrap=[1], autofilter=True, freeze_panes=True):

    workbook = xlsxwriter.Workbook(filename)
    worksheet = workbook.add_worksheet()

    for i in range(len(data.columns)):
        worksheet.write(0, i, data.columns[i])

    for i in range(len(data)):
        for j in range(len(data.columns)):
            if pd.isna(data.iloc[i, j]):
                pass
            else:
                worksheet.write(i+1, j, data.iloc[i, j])

    worksheet.set_column(0, len(data.columns), width=column_width)

    for i in range(len(data)):
        a = {}
        if i in list(row_colours.keys()):
            a["bg_color"] = row_colours[i]
        if i in row_bold:
            a["bold"] = True
        if i in row_wrap:
            a["text_wrap"] = True
        if len(a) != 0:
            worksheet.set_row(i, cell_format=workbook.add_format(a))

    if autofilter:
        worksheet.autofilter(1, 0, len(data), len(data.columns)-1)

    if freeze_panes:
        worksheet.freeze_panes(2, 0)

    workbook.close()


class process_questionnaire():
    def __init__(self):
        self.attachments = None
        self.survey = None
        self.choices = None
        self.settings = None
        self.form_title = None
        self.form_version = None

    def get_data_from_files(self, form_filename, attachements_list_filenames):
        survey = pd.read_excel(form_filename, na_values=[
                               ' ', ''], keep_default_na=False, sheet_name="survey").dropna(how='all')
        self.survey = survey
        choices = pd.read_excel(form_filename, sheet_name="choices", na_values=[
                                ' ', ''], keep_default_na=False).dropna(how='all')
        self.choices = choices
        settings = pd.read_excel(form_filename, sheet_name="settings", na_values=[
            ' ', ''], keep_default_na=False).dropna(how='all')
        self.settings = settings
        self.form_title = self.settings['form_title'].iloc[0]
        self.form_version = self.settings['version'].iloc[0]
        attachments = {}
        for j in attachements_list_filenames:
            attachments[j] = pd.read_csv(j)
        self.attachments = attachments

    def get_data_from_odk_object(self, odk_object):
        self.survey = odk_object.survey
        self.choices = odk_object.choices
        self.settings = odk_object.settings
        self.form_title = self.settings['form_title'].iloc[0]
        self.form_version = self.settings['version'].iloc[0]
        self.attachments = odk_object.attachments

    def process(self):

        def set_repeat_table_header(row):
            """ set repeat table row on every new page
            """
            tr = row._tr
            trPr = tr.get_or_add_trPr()
            tblHeader = OxmlElement('w:tblHeader')
            tblHeader.set(qn('w:val'), "true")
            trPr.append(tblHeader)
            return row

        document = dcx.Document()
        section = document.sections[-1]
        new_width, new_height = section.page_height, section.page_width
        section.orientation = dcx.enum.section.WD_ORIENT.LANDSCAPE
        section.page_width = new_width
        section.page_height = new_height
        section.top_margin = dcx.shared.Cm(0.5)
        section.bottom_margin = dcx.shared.Cm(0.5)
        section.left_margin = dcx.shared.Cm(1)
        section.right_margin = dcx.shared.Cm(1)
        heading = document.add_heading(
            f"Form title = {self.form_title}\nForm version = {str(self.form_version)}\n\n")
        heading.alignment = 1
        table = document.add_table(rows=1, cols=6)
        table.style = "Table Grid"
        title_cells = table.rows[0].cells
        headings = ["Question code", "Question type",
                    "Question", "Hint", "Select Options", "Logic"]
        for i in range(len(headings)):
            paragraph = title_cells[i].paragraphs[0]
            run = paragraph.add_run(headings[i])
            run.bold = True

        def get_choices(choice):
            choice_names = list(
                self.choices["label::English (en)"].loc[self.choices["list_name"] == choice].map(str))
            choice_labels = list(
                self.choices["name"].loc[self.choices["list_name"] == choice].map(str))
            zipped = list(zip(choice_labels, choice_names))
            zipped = [" = ".join(x) for x in zipped]
            if len(zipped) > 50:
                zipped1 = zipped[0:25]
                zipped2 = zipped[-25:-1]
                zipped = zipped1+["...The list is longer than 50, some elements are omitted..."]+zipped2
            zipped = "\n".join(zipped)
            return zipped

        def get_from_file(file):
            choice_list = list(self.attachments[file]["label"].map(str))
            if len(choice_list) > 50:
                choice_list1 = choice_list[0:25]
                choice_list2 = choice_list[-25:-1]
                choice_list = choice_list1 + \
                    ["...The list is longer than 50, some elements are omitted..."]+choice_list2
            return "\n".join(choice_list)

        def process_enclosing_variables(data):
            x = findall(r"\${.+?}", data)
            for i in x:
                data = data.replace(i, i[2:-1])
            return data

        def process_current_input(data):
            x = findall(r"\.", data)
            for i in x:
                data = data.replace(i, "input ")
            return data

        def process_string_only(column, index, cell_index):
            if not pd.isna(s[column].iloc[index]):

                paragraph = row_cells[cell_index].paragraphs[0]
                run = paragraph.add_run(process_enclosing_variables(
                    str(s[column].iloc[index])))
                if self.survey["type"].iloc[index].split(" ")[0] == "calculate":
                    run.italics = True
                    run.font.size = Pt(8)
            else:
                run = paragraph.add_run("")
                if self.survey["type"].iloc[index].split(" ")[0] == "calculate":
                    run.italics = True
                    run.font.size = Pt(8)

        def process_string_only_combined(cells, index):
            columns_names = ["Skip logic: ", "Default value: ", "Constrain value: ",
                             "Calculation: ", "Choice filter: ", "Repeat count: "]
            column_labels = ["relevant", "default", "constraint",
                             "calculation", "choice_filter", "repeat_count"]
            paragraph = cells[5].paragraphs[0]
            count = 0
            for j in range(len(column_labels)):
                if not pd.isna(self.survey[column_labels[j]].iloc[index]):
                    count += 1
            for j in range(len(column_labels)):
                if not pd.isna(self.survey[column_labels[j]].iloc[index]):
                    run = paragraph.add_run(columns_names[j])
                    run.italics = True
                    run.font.color.rgb = RGBColor(50, 0, 255)
                    if self.survey["type"].iloc[index].split(" ")[0] == "calculate":
                        run.font.size = Pt(8)
                    if column_labels[j] == "constraint":
                        run = paragraph.add_run(
                            f"{process_current_input(process_enclosing_variables(str(self.survey[column_labels[j]].iloc[index])))}")
                        if self.survey["type"].iloc[index].split(" ")[0] == "calculate":
                            run.font.size = Pt(8)
                    else:
                        run = paragraph.add_run(
                            f"{process_enclosing_variables(str(self.survey[column_labels[j]].iloc[index]))}")
                        if self.survey["type"].iloc[index].split(" ")[0] == "calculate":
                            run.font.size = Pt(8)
                    count -= 1
                    if count != 0:
                        run = paragraph.add_run("\n")

        def cell_shading(index=None, index_counter=None, header_row=False):
            if index != None:
                sss = self.survey["type"].iloc[index].split(" ")[0]
                if sss == "begin_group":
                    shading_elm_0 = parse_xml(
                        r'<w:shd {} w:fill="BDBD51"/>'.format(nsdecls('w')))
                    table.rows[index+1 -
                               index_counter].cells[0]._tc.get_or_add_tcPr().append(shading_elm_0)
                    shading_elm_1 = parse_xml(
                        r'<w:shd {} w:fill="BDBD51"/>'.format(nsdecls('w')))
                    table.rows[index+1 -
                               index_counter].cells[1]._tc.get_or_add_tcPr().append(shading_elm_1)
                    shading_elm_2 = parse_xml(
                        r'<w:shd {} w:fill="BDBD51"/>'.format(nsdecls('w')))
                    table.rows[index+1 -
                               index_counter].cells[2]._tc.get_or_add_tcPr().append(shading_elm_2)
                    shading_elm_3 = parse_xml(
                        r'<w:shd {} w:fill="BDBD51"/>'.format(nsdecls('w')))
                    table.rows[index+1 -
                               index_counter].cells[3]._tc.get_or_add_tcPr().append(shading_elm_3)
                    shading_elm_4 = parse_xml(
                        r'<w:shd {} w:fill="BDBD51"/>'.format(nsdecls('w')))
                    table.rows[index+1 -
                               index_counter].cells[4]._tc.get_or_add_tcPr().append(shading_elm_4)
                    shading_elm_5 = parse_xml(
                        r'<w:shd {} w:fill="BDBD51"/>'.format(nsdecls('w')))
                    table.rows[index+1 -
                               index_counter].cells[5]._tc.get_or_add_tcPr().append(shading_elm_5)
                elif sss == "end_group":
                    shading_elm_0 = parse_xml(
                        r'<w:shd {} w:fill="E2E2B3"/>'.format(nsdecls('w')))
                    table.rows[index+1 -
                               index_counter].cells[0]._tc.get_or_add_tcPr().append(shading_elm_0)
                    shading_elm_1 = parse_xml(
                        r'<w:shd {} w:fill="E2E2B3"/>'.format(nsdecls('w')))
                    table.rows[index+1 -
                               index_counter].cells[1]._tc.get_or_add_tcPr().append(shading_elm_1)
                    shading_elm_2 = parse_xml(
                        r'<w:shd {} w:fill="E2E2B3"/>'.format(nsdecls('w')))
                    table.rows[index+1 -
                               index_counter].cells[2]._tc.get_or_add_tcPr().append(shading_elm_2)
                    shading_elm_3 = parse_xml(
                        r'<w:shd {} w:fill="E2E2B3"/>'.format(nsdecls('w')))
                    table.rows[index+1 -
                               index_counter].cells[3]._tc.get_or_add_tcPr().append(shading_elm_3)
                    shading_elm_4 = parse_xml(
                        r'<w:shd {} w:fill="E2E2B3"/>'.format(nsdecls('w')))
                    table.rows[index+1 -
                               index_counter].cells[4]._tc.get_or_add_tcPr().append(shading_elm_4)
                    shading_elm_5 = parse_xml(
                        r'<w:shd {} w:fill="E2E2B3"/>'.format(nsdecls('w')))
                    table.rows[index+1 -
                               index_counter].cells[5]._tc.get_or_add_tcPr().append(shading_elm_5)
                elif sss == "begin_repeat":
                    shading_elm_0 = parse_xml(
                        r'<w:shd {} w:fill="AAECD3"/>'.format(nsdecls('w')))
                    table.rows[index+1 -
                               index_counter].cells[0]._tc.get_or_add_tcPr().append(shading_elm_0)
                    shading_elm_1 = parse_xml(
                        r'<w:shd {} w:fill="AAECD3"/>'.format(nsdecls('w')))
                    table.rows[index+1 -
                               index_counter].cells[1]._tc.get_or_add_tcPr().append(shading_elm_1)
                    shading_elm_2 = parse_xml(
                        r'<w:shd {} w:fill="AAECD3"/>'.format(nsdecls('w')))
                    table.rows[index+1 -
                               index_counter].cells[2]._tc.get_or_add_tcPr().append(shading_elm_2)
                    shading_elm_3 = parse_xml(
                        r'<w:shd {} w:fill="AAECD3"/>'.format(nsdecls('w')))
                    table.rows[index+1 -
                               index_counter].cells[3]._tc.get_or_add_tcPr().append(shading_elm_3)
                    shading_elm_4 = parse_xml(
                        r'<w:shd {} w:fill="AAECD3"/>'.format(nsdecls('w')))
                    table.rows[index+1 -
                               index_counter].cells[4]._tc.get_or_add_tcPr().append(shading_elm_4)
                    shading_elm_5 = parse_xml(
                        r'<w:shd {} w:fill="AAECD3"/>'.format(nsdecls('w')))
                    table.rows[index+1 -
                               index_counter].cells[5]._tc.get_or_add_tcPr().append(shading_elm_5)
                elif sss == "end_repeat":
                    shading_elm_0 = parse_xml(
                        r'<w:shd {} w:fill="C5E8DA"/>'.format(nsdecls('w')))
                    table.rows[index+1 -
                               index_counter].cells[0]._tc.get_or_add_tcPr().append(shading_elm_0)
                    shading_elm_1 = parse_xml(
                        r'<w:shd {} w:fill="C5E8DA"/>'.format(nsdecls('w')))
                    table.rows[index+1 -
                               index_counter].cells[1]._tc.get_or_add_tcPr().append(shading_elm_1)
                    shading_elm_2 = parse_xml(
                        r'<w:shd {} w:fill="C5E8DA"/>'.format(nsdecls('w')))
                    table.rows[index+1 -
                               index_counter].cells[2]._tc.get_or_add_tcPr().append(shading_elm_2)
                    shading_elm_3 = parse_xml(
                        r'<w:shd {} w:fill="C5E8DA"/>'.format(nsdecls('w')))
                    table.rows[index+1 -
                               index_counter].cells[3]._tc.get_or_add_tcPr().append(shading_elm_3)
                    shading_elm_4 = parse_xml(
                        r'<w:shd {} w:fill="C5E8DA"/>'.format(nsdecls('w')))
                    table.rows[index+1 -
                               index_counter].cells[4]._tc.get_or_add_tcPr().append(shading_elm_4)
                    shading_elm_5 = parse_xml(
                        r'<w:shd {} w:fill="C5E8DA"/>'.format(nsdecls('w')))
                    table.rows[index+1 -
                               index_counter].cells[5]._tc.get_or_add_tcPr().append(shading_elm_5)
                elif (sss == "calculate"):
                    shading_elm_0 = parse_xml(
                        r'<w:shd {} w:fill="CECECE"/>'.format(nsdecls('w')))
                    table.rows[index+1 -
                               index_counter].cells[0]._tc.get_or_add_tcPr().append(shading_elm_0)
                    shading_elm_1 = parse_xml(
                        r'<w:shd {} w:fill="CECECE"/>'.format(nsdecls('w')))
                    table.rows[index+1 -
                               index_counter].cells[1]._tc.get_or_add_tcPr().append(shading_elm_1)
                    shading_elm_2 = parse_xml(
                        r'<w:shd {} w:fill="CECECE"/>'.format(nsdecls('w')))
                    table.rows[index+1 -
                               index_counter].cells[2]._tc.get_or_add_tcPr().append(shading_elm_2)
                    shading_elm_3 = parse_xml(
                        r'<w:shd {} w:fill="CECECE"/>'.format(nsdecls('w')))
                    table.rows[index+1 -
                               index_counter].cells[3]._tc.get_or_add_tcPr().append(shading_elm_3)
                    shading_elm_4 = parse_xml(
                        r'<w:shd {} w:fill="CECECE"/>'.format(nsdecls('w')))
                    table.rows[index+1 -
                               index_counter].cells[4]._tc.get_or_add_tcPr().append(shading_elm_4)
                    shading_elm_5 = parse_xml(
                        r'<w:shd {} w:fill="CECECE"/>'.format(nsdecls('w')))
                    table.rows[index+1 -
                               index_counter].cells[5]._tc.get_or_add_tcPr().append(shading_elm_5)
            if header_row:
                shading_elm_0 = parse_xml(
                    r'<w:shd {} w:fill="919191"/>'.format(nsdecls('w')))
                table.rows[0].cells[0]._tc.get_or_add_tcPr().append(
                    shading_elm_0)
                shading_elm_1 = parse_xml(
                    r'<w:shd {} w:fill="919191"/>'.format(nsdecls('w')))
                table.rows[0].cells[1]._tc.get_or_add_tcPr().append(
                    shading_elm_1)
                shading_elm_2 = parse_xml(
                    r'<w:shd {} w:fill="919191"/>'.format(nsdecls('w')))
                table.rows[0].cells[2]._tc.get_or_add_tcPr().append(
                    shading_elm_2)
                shading_elm_3 = parse_xml(
                    r'<w:shd {} w:fill="919191"/>'.format(nsdecls('w')))
                table.rows[0].cells[3]._tc.get_or_add_tcPr().append(
                    shading_elm_3)
                shading_elm_4 = parse_xml(
                    r'<w:shd {} w:fill="919191"/>'.format(nsdecls('w')))
                table.rows[0].cells[4]._tc.get_or_add_tcPr().append(
                    shading_elm_4)
                shading_elm_5 = parse_xml(
                    r'<w:shd {} w:fill="919191"/>'.format(nsdecls('w')))
                table.rows[0].cells[5]._tc.get_or_add_tcPr().append(
                    shading_elm_5)

        cell_shading(header_row=True)
        set_repeat_table_header(table.rows[0])
        index_counter = 0
        for i in range(len(self.survey)):
            s = self.survey
            c = self.choices
            if self.survey["relevant"].iloc[i] == "false()":
                index_counter += 1
            elif self.survey["type"].iloc[i].split(" ")[0] == "start":
                index_counter += 1
            elif self.survey["type"].iloc[i].split(" ")[0] == "end":
                index_counter += 1
            elif self.survey["type"].iloc[i].split(" ")[0] == "deviceid":
                index_counter += 1
            elif self.survey["type"].iloc[i].split(" ")[0] == "phonenumber":
                index_counter += 1
            else:
                row_cells = table.add_row().cells
                cell_shading(index=i, index_counter=index_counter)
                paragraph = row_cells[0].paragraphs[0]
                run = paragraph.add_run(s["name"].iloc[i] if not pd.isna(
                    s["name"].iloc[i]) else "")
                if self.survey["type"].iloc[i].split(" ")[0] == "calculate":
                    run.italics = True
                    run.font.size = Pt(8)

                paragraph = row_cells[1].paragraphs[0]
                run = paragraph.add_run(s["type"].iloc[i].split(" ")[0])
                if self.survey["type"].iloc[i].split(" ")[0] == "calculate":
                    run.italics = True
                    run.font.size = Pt(8)

                process_string_only("label::English (en)", i, 2)
                process_string_only("hint::English (en)", i, 3)

                if (s["type"].iloc[i].split(" ")[0] == "select_one") or (s["type"].iloc[i].split(" ")[0] == "select_multiple"):
                    row_cells[4].text = get_choices(
                        s["type"].iloc[i].split(" ")[1])
                elif (s["type"].iloc[i].split(" ")[0] == "select_one_from_file") or (s["type"].iloc[i].split(" ")[0] == "select_multiple_from_file"):
                    row_cells[4].text = get_from_file(
                        s["type"].iloc[i].split(" ")[1])
                else:
                    row_cells[4].text = ""
                process_string_only_combined(row_cells, i)

        document.save(f"{self.form_title}-{str(self.form_version)}.docx")
# %% #@ ODK Class

class ODK():

    def __init__(self, url):
        self.url = url
        self.form = None
        self.project = None
        self.survey = None
        self.choices = None
        self.settings = None
        self.attachments = None

    def connect(self, email, password):

        self.email = email
        self.password = password

        req = requests.post(self.url+'/v1/sessions', data=json.dumps(
            {"email": self.email, "password": self.password}), headers={'Content-Type': 'application/json'})

        self.token = req.json()["token"]
        self.headers = {'Authorization': 'Bearer '+self.token}

    def set_target(self, project_name, form_name):
        self.project_name = project_name
        self.form_name = form_name
        self.project = self.get_project()
        self.form = self.get_form()
        self.survey = self.get_survey()
        self.choiches = self.get_choices()
        self.settings = self.get_settings()
        self.attachments = self.get_attachments()

    def list_projects(self):
        req = requests.get(self.url+'/v1/projects', headers=self.headers)
        projects = [req.json()[i]["name"] for i in range(len(req.json()))]
        return projects

    def get_project(self):
        req = requests.get(self.url+'/v1/projects', headers=self.headers)
        project = [req.json()[i]["id"] for i in range(len(req.json()))
                   if req.json()[i]["name"] == self.project_name][0]

        return project

    def list_forms(self, project=None):
        req = requests.get(self.url+'/v1/projects', headers=self.headers)
        if project != None:
            project = [req.json()[i]["id"] for i in range(
                len(req.json())) if req.json()[i]["name"] == project][0]
        else:
            project = [req.json()[i]["id"] for i in range(
                len(req.json())) if req.json()[i]["name"] == self.project_name][0]
        req = requests.get(self.url+'/v1/projects/' +
                           str(project)+"/forms", headers=self.headers)
        forms = [req.json()[i]["name"] for i in range(len(req.json()))]
        return forms

    def get_form(self):

        req = requests.get(self.url+'/v1/projects/' +
                           str(self.get_project())+"/forms", headers=self.headers)
        form = [req.json()[i]["xmlFormId"] for i in range(len(req.json()))
                if req.json()[i]["name"] == self.form_name][0]

        return form

    def save_form(self, path="", save_file=True, xml=False):

        if xml:
            extension = '.xml'
        else:
            extension = '.xlsx'
            version = str(pd.read_excel(BytesIO(req),
                                        sheet_name="settings")["version"].iloc[0])

        req = requests.get(self.url+'/v1/projects/'+str(self.get_project()) +
                           "/forms/"+self.get_form()+extension, headers=self.headers).content

        if save_file:
            file = open(path+"form_v"+version+extension, "wb")
            file.write(req)
            file.close()
        else:
            return BytesIO(req)

    def get_submissions(self):

        req = (requests.get(self.url+'/v1/projects/' +
                            str(self.project)+"/forms/" +
                            self.form+"/submissions.csv?",
                            headers=self.headers))
        df = pd.read_csv(BytesIO(req.content))
        return df

    def get_survey(self):
        req = requests.get(self.url+'/v1/projects/'+str(self.project) +
                           "/forms/"+self.form+".xlsx", headers=self.headers)
        survey = pd.read_excel(BytesIO(req.content), na_values=[
                               ' ', ''], keep_default_na=False).dropna(how='all')
        self.survey = survey
        return survey

    def get_choices(self):
        req = requests.get(self.url+'/v1/projects/'+str(self.project) +
                           "/forms/"+self.form+".xlsx", headers=self.headers)
        choices = pd.read_excel(BytesIO(req.content), sheet_name="choices", na_values=[
                                ' ', ''], keep_default_na=False).dropna(how='all')
        self.choices = choices
        return choices

    def get_settings(self):
        req = requests.get(self.url+'/v1/projects/'+str(self.project) +
                           "/forms/"+self.form+".xlsx", headers=self.headers)
        settings = pd.read_excel(BytesIO(req.content), sheet_name="settings", na_values=[
                                 ' ', ''], keep_default_na=False).dropna(how='all')
        self.settings = settings
        return settings

    def get_repeats(self):

        req = (requests.get(self.url+'/v1/projects/' +
                            str(self.project)+"/forms/"+self.form +
                            "/submissions.csv.zip?attachments=false",
                            headers=self.headers))
        zipfile = zp.ZipFile(BytesIO(req.content))

        repeats = {}

        form_id = str(pd.read_excel(BytesIO(requests.get(self.url+'/v1/projects/'+str(self.project)+"/forms/"+self.form+".xlsx", headers=self.headers).content),
                                    sheet_name="settings")["form_id"].iloc[0])

        for j in self.survey["name"].loc[self.survey["type"] == "begin_repeat"]:
            repeats[form_id+"-" + j] = pd.read_csv(zipfile.open(
                form_id+"-" + j+".csv"), na_values=[' ', ''], keep_default_na=False).dropna(how='all')

        return repeats

    def get_attachments(self):

        req = (requests.get(self.url+'/v1/projects/' +
                            str(self.project)+"/forms/" +
                            self.form+"/attachments",
                            headers=self.headers))

        attachments = {}

        for j in req.json():
            attachments[j["name"]] = pd.read_csv(BytesIO((requests.get(self.url+'/v1/projects/' + str(self.project)+"/forms/"+self.form+"/attachments/"+j["name"], headers=self.headers)).content)) if j["name"].split(
                ".")[-1] == "csv" else BytesIO((requests.get(self.url+'/v1/projects/' + str(self.project)+"/forms/"+self.form+"/attachments/"+j["name"], headers=self.headers)).content)
        return attachments

    def get_media(self):

        req = requests.get(self.url+'/v1/projects/'+str(self.project) +
                           "/forms/"+self.form+".xlsx", headers=self.headers).content

        req = (requests.post(self.url+'/v1/projects/' +
                             str(self.project)+"/forms/" +
                             self.form+"/submissions.csv.zip?",
                             headers=self.headers))
        zipfile = zp.ZipFile(BytesIO(req.content))
        media = {}
        for name in zipfile.namelist():
            if (name.split('/')[0] == 'media') & (len(name) > 6):
                media[name.split('/')[-1]] = zipfile.read(name)
        return media

    def processing_submission(self, process_datetimes=False):

        df = self.get_submissions()

        def remove_tail(list_in):
            a = []
            for j in list_in:
                if j[-2:] == ".0":
                    a.append(j[:-2])
                else:
                    a.append(j)
            return a

        def select_one(select, value):
            x = self.survey["type"].loc[self.survey["name"] == select].iloc[0].split(" ")[
                1]
            y = self.choices["label::English (en)"].loc[self.choices["list_name"].map(lambda x: x.strip())
                                                        == x].loc[self.choices["name"] == value].iloc[0]
            return y

        def select_multiple(select, value):
            x = self.survey["type"].loc[self.survey["name"] == select].iloc[0].split(" ")[
                1]
            y = self.choices.loc[self.choices["list_name"].map(
                lambda x: x.strip()) == x]
            z = []
            for i in range(len(y)):
                if str(y["name"].iloc[i]) in remove_tail(list(str(value).split(" "))):
                    z.append(y["label::English (en)"].iloc[i])
            return " \n".join(z)

        def select_one_from_file(select, value):
            x = self.survey["type"].loc[self.survey["name"] == select].iloc[0].split(" ")[
                1]
            y = pd.read_csv(x)
            z = y["label"].loc[y["name"] == value].iloc[0]
            return z

        def select_multiple_from_file(select, value):
            x = self.survey["type"].loc[self.survey["name"] == select].iloc[0].split(" ")[
                1]
            y = pd.read_csv(x)
            z = []
            for i in range(len(y)):
                if str(y["name"].iloc[i]) in remove_tail(list(str(value).split(" "))):
                    z.append(y["label"].iloc[i])
            return " \n".join(z)

        func = {"select_one_from_file": select_one_from_file,
                "select_one": select_one, "select_multiple": select_multiple, "select_multiple_from_file": select_multiple_from_file}

        group_names = list(
            self.survey["name"].loc[self.survey["type"] == "begin_group"])
        group_names = sorted(group_names, key=len, reverse=True)

        column_names = sorted(list(set(self.survey["name"].loc[((self.survey["type"] != "begin_group") & (self.survey["type"] != "end_group") & (
            self.survey["type"] != "begin_repeat") & (self.survey["type"] != "end_repeat"))]).difference(set([np.nan, ""]))), key=len, reverse=True)

        df_columns = sorted(list(df.columns), key=len, reverse=True)

        for i in df_columns:
            for j in column_names:
                if i.endswith(j):
                    df = df.rename(columns={i: j})

        for i in df_columns:
            for j in group_names:
                if i.startswith(j):
                    df = df.rename(columns={i: i[len(j):]})

        for i in df.columns:
            # try:
            a = i
            b = self.survey["type"].loc[self.survey["name"] == a]
            if len(b) == 0:
                pass
            else:
                b = b.iloc[0].split(" ")[0]
                if b in list(func.keys()):
                    for j in range(len(df)):
                        if pd.isna(df[i].iloc[j]):
                            pass
                        else:
                            try:
                                df[i].iat[j] = func[b](a, df[i].iat[j])
                            except:
                                pass

        df = df.loc[df["ReviewState"] != "rejected"]

        if process_datetimes:
            df["SubmissionDate"] = pd.to_datetime(
                df["SubmissionDate"], format="%Y-%m-%dT%H:%M:%S.%fZ")
            df["start"] = pd.to_datetime(
                df["start"], format="%Y-%m-%dT%H:%M:%S.%f%z")

            for j in self.survey["name"].loc[self.survey["type"] == "datetime"]:
                try:
                    df[j] = pd.to_datetime(
                        df[j], format="%Y-%m-%dT%H:%M:%S.%f%z")
                except:
                    df[j] = pd.to_datetime(df[j], format="mixed")

            for j in self.survey["name"].loc[self.survey["type"] == "date"]:
                try:
                    df[j] = pd.to_datetime(df[j], format="%Y-%m-%d").dt.date
                except:
                    pass

            for j in self.survey["name"].loc[self.survey["type"] == "time"]:
                try:
                    df[j] = pd.to_datetime(
                        df[j], format="%H:%M:%S.%f%z").dt.time
                except:
                    pass

        return df

    def processing_repeats(self, data=None, process_datetimes=False):

        repeats = self.get_repeats()
        df = self.processing_submission() if type(data) == type(None) else data
        set_not_rejected = list(df["KEY"])

        def remove_tail(list_in):
            a = []
            for j in list_in:
                if j[-2:] == ".0":
                    a.append(j[:-2])
                else:
                    a.append(j)
            return a

        def select_one(select, value):
            x = self.survey["type"].loc[self.survey["name"] == select].iloc[0].split(" ")[
                1]
            y = self.choices["label::English (en)"].loc[self.choices["list_name"]
                                                        == x].loc[self.choices["name"] == value].iloc[0]
            return y

        def select_multiple(select, value):
            x = self.survey["type"].loc[self.survey["name"] == select].iloc[0].split(" ")[
                1]
            y = self.choices.loc[self.choices["list_name"] == x]
            z = []
            for i in range(len(y)):
                if str(y["name"].iloc[i]) in remove_tail(list(str(value).split(" "))):
                    z.append(y["label::English (en)"].iloc[i])
            return " \n".join(z)

        def select_one_from_file(select, value):
            x = self.survey["type"].loc[self.survey["name"] == select].iloc[0].split(" ")[
                1]
            y = pd.read_csv(x)
            z = y["label"].loc[y["name"] == value].iloc[0]
            return z

        func = {"select_one_from_file": select_one_from_file,
                "select_one": select_one, "select_multiple": select_multiple}

        group_names = list(
            self.survey["name"].loc[self.survey["type"] == "begin_group"])
        group_names = sorted(group_names, key=len, reverse=True)

        column_names = sorted(list(set(self.survey["name"].loc[((self.survey["type"] != "begin_group") & (self.survey["type"] != "end_group") & (
            self.survey["type"] != "begin_repeat") & (self.survey["type"] != "end_repeat"))]).difference(set([np.nan]))), key=len, reverse=True)

        for k in repeats.keys():
            for i in repeats[k].columns:
                # try:
                a = i
                b = self.survey["type"].loc[self.survey["name"] == a]
                if len(b) == 0:
                    pass
                else:
                    b = b.iloc[0].split(" ")[0]
                    if b in list(func.keys()):
                        for j in range(len(repeats[k])):
                            if pd.isna(repeats[k][i].iloc[j]):
                                pass
                            else:
                                try:
                                    repeats[k][i].iat[j] = func[b](
                                        a, repeats[k][i].iat[j])
                                except:
                                    pass

        for j in repeats.keys():

            repeats[j] = repeats[j].loc[[True if repeats[j]["PARENT_KEY"].iloc[i].split(
                "/")[0] in set_not_rejected else False for i in range(len(repeats[j]))]]

            if process_datetimes:

                for i in self.survey["name"].loc[self.survey["type"] == "datetime"]:
                    if i in repeats[j].columns:
                        try:
                            repeats[j][i] = pd.to_datetime(
                                repeats[j][i], format="%Y-%m-%dT%H:%M:%S.%f%z")
                        except:
                            repeats[j][i] = pd.to_datetime(
                                repeats[j][i], format="mixed")

                for i in self.survey["name"].loc[self.survey["type"] == "date"]:
                    if i in repeats[j].columns:
                        try:
                            repeats[j][i] = pd.to_datetime(
                                repeats[j][i], format="%Y-%m-%d").dt.date
                        except:
                            repeats[j][i] = pd.to_datetime(
                                repeats[j][i], format="mixed").dt.date

                for i in self.survey["name"].loc[self.survey["type"] == "time"]:
                    if i in repeats[j].columns:
                        try:
                            repeats[j][i] = pd.to_datetime(
                                repeats[j][i], format="%H:%M:%S.%f%z").dt.time
                        except:
                            repeats[j][i] = pd.to_datetime(
                                repeats[j][i], format="mixed").dt.time

        return repeats

    def process_all(self, variable='', time_variable='start', process_datetimes=False):

        submissions = self.processing_submission(
            process_datetimes=process_datetimes)
        survey = self.survey.dropna(how='all')
        choices = self.choices
        settings = self.settings
        repeats = self.processing_repeats(process_datetimes=process_datetimes)
        survey_name = self.form_name
        variable = variable
        time_variable = time_variable
        media = self.get_media()
        attachments = self.attachments

        return Form(submissions, survey, choices, settings, repeats, survey_name, variable, time_variable, media, attachments)

    def save_main(self, data=None, path=""):

        df = self.processing_submission() if type(data) == type(None) else data

        df_out = copy.deepcopy(df)

        for j in df_out.select_dtypes(include=['datetime64', 'datetimetz']).columns:
            df_out[j] = df_out[j].astype(str)
        if 'start' in df_out.columns:
            df_out['start'] = df_out['start'].astype(str)
        a = []
        for j in df.columns:
            if j in list(self.survey["name"]):
                x = self.survey["label::English (en)"].loc[self.survey["name"]
                                                           == j].iloc[0]
                a.append(x)
            else:
                a.append(np.nan)

        df_out.loc[-1] = a

        df_out.sort_index(inplace=True)

        save_to_excel(df_out, path+self.form_name+"_submissions.xlsx")

    def save_repeat(self, data=None, path=""):

        repeats = self.processing_repeats() if type(data) == type(None) else data

        for k in repeats.keys():

            rep_out = copy.deepcopy(repeats[k])

            for j in rep_out.select_dtypes(include=['datetime64', 'datetimetz']).columns:
                rep_out[j] = rep_out[j].astype(str)

            a = []
            for j in repeats[k].columns:
                if j in list(self.survey["name"]):
                    x = self.survey["label::English (en)"].loc[self.survey["name"]
                                                               == j].iloc[0]
                    a.append(x)
                else:
                    a.append(np.nan)

            rep_out.loc[-1] = a

            rep_out.sort_index(inplace=True)

            save_to_excel(rep_out, path+k+".xlsx")

    def save_data(self, path=""):

        req = requests.get(self.url+'/v1/projects/'+str(self.project) +
                           "/forms/"+self.form+".xlsx", headers=self.headers).content

        version = str(pd.read_excel(BytesIO(req),
                                    sheet_name="settings")["version"].iloc[0])
        req = (requests.post(self.url+'/v1/projects/' +
                             str(self.project)+"/forms/" +
                             self.form+"/submissions.csv.zip?",
                             headers=self.headers))

        file = open(path+self.form_name+"_v"+version+".zip", "wb")
        file.write(req.content)
        file.close()

    def listing_submissions(self):

        req = (requests.get(self.url+'/v1/projects/' +
                            str(self.project)+"/forms/" +
                            self.form+"/submissions",
                            headers=self.headers))
        return req.json()

    def get_submission_metadata(self, instance):

        req = (requests.get(self.url+'/v1/projects/' +
                            str(self.project)+"/forms/" +
                            self.form+"/submissions/"+instance,
                            headers=self.headers))
        return req.json()

    def get_submission_xml(self, instance):

        req = (requests.get(self.url+'/v1/projects/' +
                            str(self.project)+"/forms/" +
                            self.form+"/submissions/"+instance+".xml",
                            headers=self.headers))
        return req.content

    def put_submission(self, instance, data):

        req = (requests.put(url=self.url+'/v1/projects/' +
                            str(self.project)+"/forms/" +
                            self.form+"/submissions/"+instance, data=data,
                            headers=self.headers))
        return req

    def get_parent_tag(self, tag):

        n = self.survey.loc[self.survey['name'] == tag].index[0]
        begin_group = len(
            self.survey.iloc[:n].loc[self.survey['type'] == 'begin_group'])
        end_group = len(
            self.survey.iloc[:n].loc[self.survey['type'] == 'end_group'])
        begin_repeat = len(
            self.survey.iloc[:n].loc[self.survey['type'] == 'begin_repeat'])
        end_repeat = len(
            self.survey.iloc[:n].loc[self.survey['type'] == 'end_repeat'])

        if end_repeat < begin_repeat:
            return (self.survey['name'].iloc[:n].loc[self.survey['type'] == 'begin_repeat']).iloc[-1]
        elif end_group < begin_group:
            return (self.survey['name'].iloc[:n].loc[self.survey['type'] == 'begin_group']).iloc[-1]
        else:
            return None

    def return_element(self, tree, data: str):
        for elem in tree.iter():
            if elem.tag == data:
                return elem
            else:
                pass
        return None

    def modify_variable_xml(self, xml, variable: str, function):
        tree = ET.parse(BytesIO(xml))
        d = self.return_element(tree, variable)
        if d == None:
            print(f"{variable} is not in the xml")
            return xml
        else:
            try:
                k = d.text
                d.text = function(k)
                xml_out = BytesIO()
                tree.write(xml_out, encoding='utf-8')
                return xml_out.getvalue()
            except:
                print('an error occurred while processing for variable ', variable)
                return xml

    def update_xml(self, xml):

        tree = ET.parse(BytesIO(xml))
        root = tree.getroot()

        if tree.find('meta').find('deprecatedID') == None:
            old = tree.find('meta').find('instanceID').text
            tree.find('meta').find(
                'instanceID').text = 'uuid:'+str(uuid.uuid4())
            deprecated = ET.Element("deprecatedID")
            deprecated.text = old
            root.find('meta').append(deprecated)

        else:
            if len(tree.find('meta').find('deprecatedID').text) > 0:
                old = tree.find('meta').find('instanceID').text
                tree.find('meta').find(
                    'instanceID').text = 'uuid:'+str(uuid.uuid4())
                root.find('meta').find('deprecatedID').text = old
        xml_out = BytesIO()
        tree.write(xml_out, encoding='utf-8')
        return xml_out.getvalue()

    def change_submission(self, xml, id):
        self.put_submission(id, self.update_xml(xml))

    def drop_variable_xml(self, xml, variable: str, parent_tag=None):

        tree = ET.parse(BytesIO(xml))
        root = tree.getroot()
        for elem in tree.iter():
            if elem.tag == variable:
                if parent_tag == None:
                    root.remove(elem)
                else:
                    self.return_element(tree, parent_tag).remove(elem)
        xml_out = BytesIO()
        tree.write(xml_out, encoding='utf-8')
        return xml_out.getvalue()

    def add_variable_xml(self, xml, variable: str, parent_tag=None):

        tree = ET.parse(BytesIO(xml))
        root = tree.getroot()
        if type(self.return_element(tree, variable)) == type(None):
            if parent_tag == None:
                child = ET.SubElement(root, variable)
            else:
                child = ET.SubElement(
                    self.return_element(tree, parent_tag), variable)
            xml_out = BytesIO()
            tree.write(xml_out, encoding='utf-8')
            return xml_out.getvalue()
