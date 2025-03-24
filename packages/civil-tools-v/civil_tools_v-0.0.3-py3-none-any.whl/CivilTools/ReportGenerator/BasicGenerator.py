import warnings
from enum import Enum
from docx import Document
from docx.document import Document as doc
from docx.shared import Inches, RGBColor, Pt, Cm
from docx.oxml.ns import qn, nsdecls
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml.parser import parse_xml
from docx.oxml.xmlchemy import BaseOxmlElement
from docx.enum.section import WD_ORIENTATION
from .DocParagraph import DocParagraph
from .DocPicture import DocPicture
from .DocTable import DocTable
from .UtilFunctions import set_cell_border, analysis_sub_and_super_script

FONT_STR = "w:eastAsia"


class Page:
    def __init__(self, width: float, height: float, is_landscape=False):
        self._width = width
        self._height = height
        self.is_landscape = is_landscape

    @property
    def width(self):
        return self._height if self.is_landscape else self._width

    @property
    def height(self):
        return self._width if self.is_landscape else self._height


class PageSize(Enum):
    A4 = Page(210, 297)
    A4_LANDSCAPE = Page(210, 297, True)
    A3 = Page(297, 420)
    A3_LANDSCAPE = Page(297, 420, True)


class BasicGenerator:

    def __init__(self):
        self.mm_factor = 36000
        self.doc = Document()
        self.doc: doc
        self.init_styles()
        self.highlighted_color = 7

    def init_styles(self):
        styles = self.doc.styles
        self.body_style = self.init_body_style(styles)
        self.table_style = self.init_table_style(styles)
        self.small_title_style = self.init_small_title_style(styles)

    def init_body_style(self, styles):
        """生成一个用于普通段落文字的body_style

        Args:
            styles (_type_): _description_

        Returns:
            _type_: _description_
        """
        body_style = styles["Body Text"]
        body_style.font.name = "Times New Roman"
        body_style.paragraph_format.first_line_indent = Inches(0.32)
        body_style.paragraph_format.space_before = Pt(0)
        body_style.paragraph_format.space_after = Pt(0)
        body_style._element.rPr.rFonts.set(qn(FONT_STR), "宋体")
        return body_style

    def init_table_style(self, styles):
        """生成一个用于表格内部段落文字的table_style

        Args:
            styles (_type_): _description_

        Returns:
            _type_: _description_
        """
        table_style = styles["Normal"]
        table_style.font.name = "Time New Roman"
        table_style.font.size = Pt(10)
        table_style._element.rPr.rFonts.set(qn(FONT_STR), "宋体")
        table_style.paragraph_format.space_before = Pt(0)
        table_style.paragraph_format.space_after = Pt(0)
        return table_style

    def init_small_title_style(self, styles):
        """生成用于表格名称和图片名称的small_title_style

        Args:
            styles (_type_): _description_

        Returns:
            _type_: _description_
        """
        small_title_style = styles["Body Text 2"]
        small_title_style.font.name = "Times New Roman"
        small_title_style._element.rPr.rFonts.set(qn(FONT_STR), "黑体")
        small_title_style.paragraph_format.space_before = Pt(6)
        small_title_style.paragraph_format.space_after = Pt(3)
        small_title_style.paragraph_format.line_spacing = Pt(15)
        return small_title_style

    def change_paper_size(self, page_size: PageSize, column_num: int = 1):
        section = self.doc.sections[0]
        section.page_width = int(page_size.value.width * self.mm_factor)
        section.page_heigth = int(page_size.value.height * self.mm_factor)
        if page_size.value.is_landscape:
            section.orientation = WD_ORIENTATION.LANDSCAPE
        sect_pr = section._sectPr
        cols = sect_pr.xpath("./w:cols")[0]
        cols.set(qn("w:num"), str(column_num))

    def change_paper_margin(self, left: int, top: int, right: int, bottom: int):
        """改变页边距，单位为mm，按照左、上、右、下的顺序"""
        section = self.doc.sections[0]
        section.top_margin = int(self.mm_factor * top)
        section.bottom_margin = int(self.mm_factor * bottom)
        section.left_margin = int(self.mm_factor * left)
        section.right_margin = int(self.mm_factor * right)

    def add_big_title(self, context):
        p = self.doc.add_paragraph(context)
        run = p.runs[0]
        run.font.size = Pt(14)
        run.font.bold = True
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    def add_title(
        self, title: DocParagraph, space_before: float = 0, space_after: float = 0
    ):
        """为了添加大纲等级而添加的函数，不需要大纲等级时不要使用

        Args:
            title (DocParagraph): _description_
        """
        p = self.doc.add_heading("", level=title.par_level)
        run = p.add_run(title.context)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.color.rgb = RGBColor(0, 0, 0)
        par_format = p.paragraph_format
        par_format.space_before = Pt(space_before)
        par_format.space_after = Pt(space_after)

        self.paragraph_format_double_check(p, title)

    def add_paragraph(self, par: DocParagraph):
        if par.context == None or par.context == "":
            warnings.warn("This par has no context.")
            return
        if "{" in par.context:
            a, b = analysis_sub_and_super_script(par.context)
            if par.style != None:
                p = self.doc.add_paragraph("", par.style)
            else:
                p = self.doc.add_paragraph()
            self.add_context_with_sub_or_super(p, a, b)
        elif par.style != None:
            p = self.doc.add_paragraph(par.context, par.style)
        else:
            p = self.doc.add_paragraph(par.context)
        self.paragraph_format_double_check(p, par)

    def add_context_with_sub_or_super(self, p, str_list, sub_or_super):
        for i in range(len(str_list)):
            run = p.add_run(str_list[i])
            run.font.name = "Times New Roman"
            if sub_or_super[i] == 1:
                run.font.subscript = True
            elif sub_or_super[i] == 2:
                run.font.superscript = True
            elif sub_or_super[i] == 3:
                run.font.highlight_color = self.highlighted_color
            run.font.size = Pt(11)

    def paragraph_format_double_check(self, p, doc_par: DocParagraph):
        run = p.runs[0]
        if doc_par.font_size != None:
            run.font.size = Pt(doc_par.font_size)
        if doc_par.is_bold != None:
            run.font.bold = doc_par.is_bold
        par_format = p.paragraph_format
        if doc_par.first_line_indent != None:
            par_format.first_line_indent = Inches(doc_par.first_line_indent)
        if doc_par.alignment != None:
            par_format.alignment = doc_par.alignment
        if doc_par.font_size != None:
            run.font.size = Pt(doc_par.font_size)

    def add_table(self, my_table: DocTable):
        if my_table.title != None:
            p = self.doc.add_paragraph(my_table.title, self.small_title_style)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        table = self.doc.add_table(
            rows=my_table.row_num, cols=my_table.column_num, style="Table Grid"
        )
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, j, k, p in my_table.merged_cells:
            table.cell(i, j).merge(table.cell(k, p))
        for i in range(my_table.row_num):
            for j in range(my_table.column_num):
                self.add_table_cell(table, my_table, i, j)
        self.tab_bg_color(table, my_table.row_num, my_table.column_num)

    def tab_bg_color(self, table, rows, cols, color_str="CCCCCC"):
        """表格样式调整"""
        shading_list = locals()
        for i in range(cols):
            shading_list["shading_elm_" + str(i)] = parse_xml(
                r'<w:shd {} w:fill="{bgColor}"/>'.format(
                    nsdecls("w"), bgColor=color_str
                )
            )
            table.rows[0].cells[i]._tc.get_or_add_tcPr().append(
                shading_list["shading_elm_" + str(i)]
            )
        for i in range(rows):
            shading_list["shading_elm_" + str(cols + i)] = parse_xml(
                r'<w:shd {} w:fill="{bgColor}"/>'.format(
                    nsdecls("w"), bgColor=color_str
                )
            )
            table.rows[i].cells[0]._tc.get_or_add_tcPr().append(
                shading_list["shading_elm_" + str(cols + i)]
            )
        table_cells = table._cells
        for i in range(rows):
            table_row = table_cells[i * cols : i * cols + cols]
            for k in range(cols):
                table_row[k].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                table_row[k].paragraphs[
                    0
                ].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                if i == 0:
                    runs = table_row[k].paragraphs[0].runs
                    for m in range(len(runs)):
                        runs[m].bold = True
                table.rows[i].height = Cm(0.7)

    def add_table_cell(self, table, my_table: DocTable, i, j):
        try:
            cell = table.cell(i, j)
            _ = my_table.context[i][j]
        except IndexError:
            return
        if "{" in my_table.context[i][j]:
            a, b = analysis_sub_and_super_script(my_table.context[i][j])
            p = cell.paragraphs[0]
            self.add_context_with_sub_or_super(p, a, b)
        else:
            cell.text = my_table.context[i][j]
            cell.paragraphs[0].style = self.table_style
            p = cell.paragraphs[0]
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        if my_table.no_grid:
            set_cell_border(
                cell,
                top={"color": "#FFFFFF"},
                bottom={"color": "#FFFFFF"},
                start={"color": "#FFFFFF"},
                end={"color": "#FFFFFF"},
            )
        if my_table.all_bold:
            p.runs[0].font.bold = True

    def add_picture(self, doc_picture: DocPicture):
        self.doc.add_picture(
            doc_picture.path_or_stream, width=Inches(doc_picture.width)
        )
        last_paragraph = self.doc.paragraphs[-1]
        last_paragraph.alignment = 1

    def add_blank_paragraph(self):
        self.doc.add_paragraph()

    def add_page_break(self):
        if len(self.doc.paragraphs[-1].runs) > 0:
            self.doc.paragraphs[-1].runs[0].add_break(WD_BREAK.PAGE)
        else:
            run = self.doc.paragraphs[-1].add_run("")
            run.add_break(WD_BREAK.PAGE)

    def save(self, path: str):
        for _ in range(10):
            try:
                self.doc.save(path)
            except Exception:
                path = path.replace(".docx", "1.docx")
