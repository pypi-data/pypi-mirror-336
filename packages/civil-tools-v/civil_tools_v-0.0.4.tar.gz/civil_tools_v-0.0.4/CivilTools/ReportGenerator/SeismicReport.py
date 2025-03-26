from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.shared import Inches, RGBColor, Pt, Cm

from .BasicGenerator import BasicGenerator, PageSize
from .DocParagraph import DocParagraph
from .DocTable import DocTable
from .DocPicture import DocPicture
from .SeismicReportTemplate import SRTemplate
from .UtilFunctions import add_comma_in_num_str

from ..YDBLoader.BuildingDefine import MassResult, Period
from ..FigureGenerator.BasicPltPlotter import BasicPltPlotter, SeismicPlotter


class SeismicReportData:
    def __init__(self, name: str | None = None):
        self.project_name = name
        self.floor_num = 8
        self.yjk_version = None
        self.mass_result = None
        self.period = None
        self.__mock_data()

    def __mock_data(self):
        self.yjk_version = "6.0.0"
        self.mass_result = MassResult.mock_data()
        self.period = Period.mock_data(num=11, mass_participate=0.01)

    @property
    def is_valid(self):
        return True


class SeismicReport(BasicGenerator):
    G = 9.8
    """重力加速度"""

    def __init__(self, all_data: SeismicReportData | None = None):
        super().__init__()
        self.all_data = all_data
        # 修改为A3图纸，横向，两栏
        self.change_paper_size(PageSize.A3_LANDSCAPE, 2)
        # 修改纸张Margin，单位mm
        self.change_paper_margin(32, 25, 32, 25)
        # 格式统一修改
        self.body_style.paragraph_format.line_spacing = Pt(22)

    def creat_doc(self):
        if self.all_data == None or not self.all_data.is_valid:
            raise ValueError(
                "The data is not ready, please use set_data() to assign data."
            )
        self.__add_info()
        self.__add_seismic_chapter()

    def __add_info(self):
        model_name = self.all_data.project_name
        par_context = SRTemplate.FIRST_INFO(model_name)
        paragraph = DocParagraph(par_context)
        paragraph.style = self.body_style
        self.add_paragraph(paragraph)

    def __add_seismic_chapter(self):
        chapter_index = 8
        sub_index = 1
        self.__add_seismic_chapter_title(chapter_index)
        sub_index = self.__add_seismic_embedding(chapter_index, sub_index)
        sub_index = self.__add_project_mass(chapter_index, sub_index)
        sub_index = self.__add_period(chapter_index, sub_index)
        sub_index = self.__add_shear_mass_ratio(chapter_index, sub_index)
        sub_index = self.__add_shear_and_moment(chapter_index, sub_index)
        sub_index = self.__add_horizental_moment_ratio_for_column(
            chapter_index, sub_index
        )
        sub_index = self.__add_disp_and_drift(chapter_index, sub_index)
        sub_index = self.__add_horizental_stiffness_ratio(chapter_index, sub_index)
        sub_index = self.__add_rotation_ratio(chapter_index, sub_index)
        sub_index = self.__add_stiffness_mass_ratio(chapter_index, sub_index)
        sub_index = self.__add_shear_capacity_ratio(chapter_index, sub_index)
        sub_index = self.__add_wind_acc(chapter_index, sub_index)

    def __add_seismic_chapter_title(self, chapter_index: int):
        # 获取需要的数据
        yjk_version = self.all_data.yjk_version
        # 开始生成报告
        current_context = SRTemplate.SEISMIC_CHAPTER_TITLE
        par_context = DocParagraph(current_context.title(chapter_index))
        par_context.par_level = 1
        self.add_title(par_context, 12, 6)
        paragraph_texts = current_context.paragraph(chapter_index, yjk_version)
        for context in paragraph_texts[:-1]:
            paragraph = DocParagraph(context)
            paragraph.style = self.body_style
            self.add_paragraph(paragraph)

        text = paragraph_texts[-1]
        paragraph = DocParagraph(text)
        paragraph.style = self.body_style
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        paragraph.first_line_indent = 0
        self.add_paragraph(paragraph)

        figure_title = current_context.picture(chapter_index)
        self.__insert_table_figure_title(figure_title)

    def __add_seismic_embedding(self, chapter_index: int, sub_index: int):

        current_context = SRTemplate.SEISMIC_EMBEDDING
        self.__insert_title_par_2(current_context, chapter_index, sub_index)

        context = current_context.paragraph(chapter_index, sub_index)[0]
        self.__insert_normal_para(context)

        table_title = current_context.table(chapter_index, sub_index)
        self.__insert_table_figure_title(table_title)

        table = DocTable(3, 7)
        table.merge_cells(1, 4, 2, 4)
        table.merge_cells(1, 5, 2, 5)
        table.set_table_context(current_context.table_context)
        self.add_table(table)

        return sub_index + 1

    def __add_project_mass(self, chapter_index: int, sub_index: int):
        mass_result = self.all_data.mass_result
        dead_load = mass_result.total_dead_load
        live_load = mass_result.total_live_load * 0.5
        total_load = mass_result.total_load
        total_area = mass_result.total_slab_area
        average_dead_load = f"{dead_load / total_area:.1f}"
        average_live_load = f"{live_load / total_area:.1f}"
        average_total_load = f"{total_load / total_area:.1f}"

        current_context = SRTemplate.PROJECT_MASS
        self.__insert_title_par_2(current_context, chapter_index, sub_index)

        contexts = current_context.paragraph(
            chapter_index,
            sub_index,
            total_mass=add_comma_in_num_str(int(total_load / SeismicReport.G)),
            total_area=add_comma_in_num_str(int(total_area)),
            average_load=average_total_load,
        )
        paragraph = DocParagraph(contexts[0])
        paragraph.style = self.body_style
        self.add_paragraph(paragraph)

        table_title = current_context.table(chapter_index, sub_index)
        self.__insert_table_figure_title(table_title)

        table = DocTable(4, 4)
        table.set_table_context(
            current_context.table_context(
                dead_mass=add_comma_in_num_str(int(dead_load / SeismicReport.G)),
                live_mass=add_comma_in_num_str(int(live_load / SeismicReport.G)),
                total_mass=add_comma_in_num_str(int(total_load / SeismicReport.G)),
                dead_percentage=f"{dead_load/total_load*100:.1f}%",
                live_percentage=f"{live_load/total_load*100:.1f}%",
                total_percentage="100%",
                dead_average=average_dead_load,
                live_average=average_live_load,
                total_average=average_total_load,
            )
        )
        self.add_table(table)

        return sub_index + 1

    def __add_period(self, chapter_index: int, sub_index: int):
        current_context = SRTemplate.PERIOD
        self.__insert_title_par_2(current_context, chapter_index, sub_index)

        paras = current_context.paragraph(
            self.all_data.period, chapter_index, sub_index
        )
        self.__insert_normal_para(paras[0])

        table_title = current_context.table(chapter_index, sub_index)
        self.__insert_table_figure_title(table_title)
        period_num = len(self.all_data.period.periods)
        if period_num <= 10:
            row_num = period_num + 1
        else:
            row_num = 12
        table = DocTable(row_num, 6)
        table_context = current_context.table_context
        last_mass_participate_x = 0
        last_mass_participate_y = 0
        for i in range(row_num - 1):
            temp_period = self.all_data.period.periods[i if i <= 9 else row_num - 2]
            if i <= 8 or (i == 9 and row_num == 11) or i > 9:
                table_context.append(
                    [
                        str(i + 1),
                        temp_period.time_str,
                        temp_period.movement_coeff,
                        temp_period.rotation_coeff,
                        temp_period.get_mass_participate_x(last_mass_participate_x),
                        temp_period.get_mass_participate_y(last_mass_participate_y),
                    ]
                )
            elif i == 9:
                table_context.append(["..."] * 6)
            last_mass_participate_x += temp_period.mass_participate_x
            last_mass_participate_y += temp_period.mass_participate_y
        table.set_table_context(table_context)
        self.add_table(table)

        self.__insert_normal_para(paras[1])

        text = "*{（这里需要振型图片！）}"
        titles = ["(a) 第一振型", "(b) 第二振型", "(c) 第三振型"]
        for i in range(3):
            paragraph = DocParagraph(text)
            paragraph.style = self.body_style
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            paragraph.first_line_indent = 0
            self.add_paragraph(paragraph)
            self.__insert_table_figure_title(titles[i])
        figure_title = current_context.picture(chapter_index, sub_index)
        self.__insert_table_figure_title(figure_title)

        return sub_index + 1

    def __add_shear_mass_ratio(self, chapter_index: int, sub_index: int):
        current_context = SRTemplate.SHEAR_MASS_RATIO
        self.__insert_title_par_2(current_context, chapter_index, sub_index)

        para = current_context.paragraph(chapter_index, sub_index)
        self.__insert_normal_para(para)

        table_title = current_context.table(chapter_index, sub_index)
        self.__insert_table_figure_title(table_title)
        table = DocTable(3, 4)
        table.merge_cells(1, 2, 2, 2)
        table.set_table_context(current_context.table_context())
        self.add_table(table)

        figure = SeismicPlotter()
        figure.test_plot()
        stream = figure.save_to_stream()
        picture = DocPicture(stream, 2)
        self.add_picture(picture)

        figure_title = current_context.picture(chapter_index, sub_index)
        self.__insert_table_figure_title(figure_title)

        return sub_index + 1

    def __add_shear_and_moment(self, chapter_index: int, sub_index: int):
        current_context = SRTemplate.SHEAR_AND_MOMENT
        self.__insert_title_par_2(current_context, chapter_index, sub_index)

        para = current_context.paragraph(chapter_index, sub_index)
        self.__insert_normal_para(para)

        table_title = current_context.table(chapter_index, sub_index)
        self.__insert_table_figure_title(table_title)
        table = DocTable(6, 4)
        table.merge_cells(0, 0, 1, 1)
        table.merge_cells(0, 2, 0, 3)
        table.merge_cells(2, 0, 3, 0)
        table.merge_cells(4, 0, 5, 0)
        table.set_table_context(current_context.table_context())
        self.add_table(table)

        figure = SeismicPlotter()
        figure.test_plot()
        stream = figure.save_to_stream()
        picture = DocPicture(stream, 2)
        self.add_picture(picture)

        figure_titles = current_context.picture(chapter_index, sub_index)
        self.__insert_table_figure_title(figure_titles[0])

        figure = SeismicPlotter()
        figure.test_plot()
        stream = figure.save_to_stream()
        picture = DocPicture(stream, 2)
        self.add_picture(picture)

        self.__insert_table_figure_title(figure_titles[1])

        return sub_index + 1

    def __add_horizental_moment_ratio_for_column(
        self, chapter_index: int, sub_index: int
    ):
        current_context = SRTemplate.HORIZENTAL_MOMENT_RATIO_FOR_COLUMN
        self.__insert_title_par_2(current_context, chapter_index, sub_index)

        paras = current_context.paragraph(chapter_index, sub_index)
        self.__insert_normal_para(paras[0])
        self.__insert_normal_para(paras[1])

        figure = SeismicPlotter()
        figure.test_plot()
        stream = figure.save_to_stream()
        picture = DocPicture(stream, 2)
        self.add_picture(picture)

        figure_title = current_context.picture(chapter_index, sub_index)
        self.__insert_table_figure_title(figure_title)

        table_titles = current_context.table(chapter_index, sub_index)
        self.__insert_table_figure_title(table_titles[0])
        row_num = self.all_data.floor_num + 1
        table = DocTable(row_num, 5)
        table_context = current_context.table_context()
        for _ in range(self.all_data.floor_num):
            table_context.append(["--"] * 5)
        table.set_table_context(table_context)
        self.add_table(table)

        self.__insert_table_figure_title(table_titles[1])
        row_num = self.all_data.floor_num + 1
        table = DocTable(row_num, 5)
        table_context = current_context.table_context()
        for _ in range(self.all_data.floor_num):
            table_context.append(["--"] * 5)
        table.set_table_context(table_context)
        self.add_table(table)

        return sub_index + 1

    def __add_disp_and_drift(self, chapter_index: int, sub_index: int):
        current_context = SRTemplate.DISP_AND_DRIFT
        self.__insert_title_par_2(current_context, chapter_index, sub_index)

        return sub_index + 1

    def __add_horizental_stiffness_ratio(self, chapter_index: int, sub_index: int):
        current_context = SRTemplate.HORIZENTAL_STIFFNESS_RATIO
        self.__insert_title_par_2(current_context, chapter_index, sub_index)

        return sub_index + 1

    def __add_rotation_ratio(self, chapter_index: int, sub_index: int):
        current_context = SRTemplate.ROTATION_RATIO
        self.__insert_title_par_2(current_context, chapter_index, sub_index)

        return sub_index + 1

    def __add_stiffness_mass_ratio(self, chapter_index: int, sub_index: int):
        current_context = SRTemplate.STIFFNESS_MASS_RATIO
        self.__insert_title_par_2(current_context, chapter_index, sub_index)

        return sub_index + 1

    def __add_shear_capacity_ratio(self, chapter_index: int, sub_index: int):
        current_context = SRTemplate.SHEAR_CAPACITY_RATIO
        self.__insert_title_par_2(current_context, chapter_index, sub_index)

        return sub_index + 1

    def __add_wind_acc(self, chapter_index: int, sub_index: int):
        current_context = SRTemplate.WIND_ACC
        self.__insert_title_par_2(current_context, chapter_index, sub_index)

        return sub_index + 1

    def __insert_title_par_2(self, current_context, chapter_index, sub_index):
        """用于生成二级子目录"""
        par_context = DocParagraph(current_context.title(chapter_index, sub_index))
        par_context.par_level = 2
        self.add_title(par_context, 6, 6)

    def __insert_normal_para(self, context):
        paragraph = DocParagraph(context)
        paragraph.style = self.body_style
        self.add_paragraph(paragraph)

    def __insert_table_figure_title(self, context):
        paragraph = DocParagraph(context)
        paragraph.style = self.small_title_style
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        self.add_paragraph(paragraph)
