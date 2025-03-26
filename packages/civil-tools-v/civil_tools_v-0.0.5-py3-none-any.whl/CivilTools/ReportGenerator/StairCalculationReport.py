from .BasicGenerator import BasicGenerator, PageSize
from docx.shared import Inches, RGBColor, Pt, Cm
from CivilTools.YDBLoader.BuildingDefine.StairPart import (
    Position,
    StairPart,
    LoadParams,
    StairLoad,
    LoadCalulateType,
)
from .UtilFunctions import MatrixSolver
from .DocTable import DocTable
from .DocPicture import DocPicture
from .DocParagraph import DocParagraph


class StairCalculationReport(BasicGenerator):
    def __init__(self):
        super().__init__()
        # 修改为A3图纸，横向，两栏
        self.change_paper_size(PageSize.A4)
        # 修改纸张Margin，单位mm
        self.change_paper_margin(20, 20, 20, 20)
        # 格式统一修改
        # self.body_style.paragraph_format.line_spacing = Pt(22)
        self.set_const()

    def set_const(self):
        self.big_title = "现浇板式普通楼梯设计"
        self.operate_code_title = "执行规范："
        self.operate_code_text_list = [
            "《混凝土结构通用规范》(GB 55008-2021)，本文简称《混凝土通用规范》；",
            "《混凝土结构设计规范》(GB 50010-2010(2015年版))，本文简称《混凝土规范》；",
            "《工程结构通用规范》(GB 55001-2021)；",
            "《建筑结构荷载规范》(GB 50009-2012)，本文简称《荷载规范》；",
        ]
        self.symbol_explain = "钢筋：d-HPB300；D-HRB335；E-HRB400；F-RRB400；G-HRB500；Q-HRBF400；R-HRB550"
        self.cut_line = "---------------------------------------------------------------------------------------------"
        self.concrete_data = {
            30: [14.3, 1.43, 2.01, 30000],  # // fc, ft, ftk, E
            35: [16.7, 1.57, 2.20, 31500],  # // fc, ft, ftk, E
            40: [19.1, 1.71, 2.39, 32500],  # // fc, ft, ftk, E
        }
        self.concrete_density = 25
        self.rebar_level = "HRB400"
        self.rebar_fy = 360
        self.indent = 0.22
        self.normal_font_size = 11

    # 数据入口！！！！
    def set_stair_data(self):
        """这个函数需要重写！！！！！是数据流入的最主要的入口"""
        position1 = Position(0, 2180, 0, 1910, 5030, 6850)
        sp1 = StairPart(position1, 13)
        sp1.left_thick = sp1.main_thick = sp1.right_thick = 130
        position2 = Position(0, 2180, 0, 1910, 5030, 6850)
        sp2 = StairPart(position2, 13)
        sp2.set_beam_offset(1, 500)
        position3 = Position(0, 2180, 0, 1910, 5030, 6850)
        sp3 = StairPart(position3, 13)
        sp3.set_beam_offset(2, 500)
        position4 = Position(0, 2180, 0, 1910, 5030, 6850)
        sp4 = StairPart(position4, 13)
        sp4.set_beam_offset(1, 500)
        sp4.set_beam_offset(2, 500)

        self.stair_list = [sp1, sp2, sp3, sp4]

    # 数据入口！！！！
    def set_calculate_info(self):
        self.load_param = LoadParams()
        self.concrete_level = 30
        self.rebar_area_adjust_coef = 1
        self.cover_thickness = 15

        self.displacement_limit = 200
        self.crack_width_limit = 0.3

    def load_calculate(self):
        left_dead = (
            self.current_stair.left_thick * self.concrete_density / 1000
            + self.load_param.append_dead_load
        )
        self.left_slab_load = StairLoad(
            left_dead, self.load_param.live_load, self.load_param
        )

        right_dead = (
            self.current_stair.right_thick * self.concrete_density / 1000
            + self.load_param.append_dead_load
        )
        self.right_slab_load = StairLoad(
            right_dead, self.load_param.live_load, self.load_param
        )

        main_dead = (
            self.current_stair.equivlent_main_slab_thick * self.concrete_density / 1000
            + self.load_param.append_dead_load
        )
        self.main_slab_load = StairLoad(
            main_dead, self.load_param.live_load, self.load_param
        )

    def do_the_math(self):
        solver = MatrixSolver(self.current_stair)
        solver.set_load(
            self.left_slab_load.q, self.main_slab_load.q, self.right_slab_load.q
        )
        self.current_stair.set_calculate_result(solver.submit_problem())

    def add_first_part(self):
        # 添加最大的标题
        self.add_big_title(self.big_title)
        # 添加一个2*6的表格，包含项目基本信息和设计人员信息
        table = DocTable(2, 6)
        table_context = [
            ["项目名称", "XX项目", "构件编号", "XXXX", "日    期", "2024/10/10"],
            ["设    计", "", "校    对", "", "审    核", ""],
        ]
        table.set_table_context(table_context)

        # table.no_grid = True
        table.all_bold = True
        self.add_table(table)
        # 添加执行规范和符号示意
        doc_par = DocParagraph(self.operate_code_title)
        doc_par.style = self.body_style
        doc_par.is_bold = True
        doc_par.font_size = 10
        doc_par.first_line_indent = 0
        self.add_paragraph(doc_par)

        for line in self.operate_code_text_list:
            doc_par = DocParagraph(line)
            doc_par.style = self.body_style
            doc_par.font_size = 10
            doc_par.first_line_indent = 0.22
            self.add_paragraph(doc_par)
        self.add_blank_paragraph()
        doc_par = DocParagraph(self.symbol_explain)
        doc_par.style = self.body_style
        doc_par.font_size = 10
        doc_par.first_line_indent = 0
        self.add_paragraph(doc_par)
        doc_par.context = self.cut_line
        self.add_paragraph(doc_par)

    def add_single_stair(self, stair: StairPart, index):
        self.current_stair = stair
        self.current_index = index

        # 添加一个标题
        doc_par = DocParagraph(
            f"{self.current_index}、{self.current_stair.stair_type}类梯板，标高区间：{self.current_stair.stair_elevation_range}"
        )
        doc_par.style = self.body_style
        doc_par.first_line_indent = 0
        doc_par.font_size = 11
        doc_par.is_bold = True
        doc_par.par_level = 1
        self.add_title(doc_par)

        # 进行必要的计算
        self.load_calculate()

        self.do_the_math()

        # 添加part1已知条件
        self.add_1_basic_info()
        self.add_2_load_and_calculate()
        self.add_3_result()
        self.add_blank_paragraph()

    def add_1_basic_info(self):
        stair = self.current_stair
        index = self.current_index
        indent = self.indent
        font_size = self.normal_font_size
        # 添加一个二级标题
        doc_par = DocParagraph(f"{index}.1 已知条件")
        doc_par.style = self.body_style
        doc_par.first_line_indent = 0
        doc_par.font_size = font_size
        doc_par.is_bold = False
        doc_par.par_level = 2
        self.add_title(doc_par)

        # 几何信息
        doc_par = DocParagraph("几何信息：")
        doc_par.style = self.body_style
        doc_par.first_line_indent = indent
        doc_par.font_size = font_size
        self.add_paragraph(doc_par)

        doc_par = DocParagraph(
            f"左标高={stair.position.left_elevation/1000:.3f}m；右标高={stair.position.right_elevation/1000:.3f}m"
        )
        doc_par.style = self.body_style
        doc_par.first_line_indent = indent * 2
        doc_par.font_size = font_size
        self.add_paragraph(doc_par)
        doc_par.context = f"左平台长度={stair.position.left_plat_length}mm；右平台长度={stair.position.right_plat_length}mm"
        self.add_paragraph(doc_par)
        doc_par.context = (
            f"左平台厚度={stair.left_thick}mm；右平台厚度={stair.right_thick}mm"
        )
        self.add_paragraph(doc_par)
        doc_par.context = f"内延长（左）={stair.beam_list[1].offset}mm；内延长（右）={stair.beam_list[2].offset}mm"
        self.add_paragraph(doc_par)
        doc_par.context = f"梯段长度={stair.position.right_x1 - stair.position.left_x2}mm；踏步数={stair.step_num}"
        doc_par.is_bold = True
        self.add_paragraph(doc_par)
        doc_par.context = f"梯板厚度={stair.main_thick}mm；梯段宽度={stair.stair_width}mm；梯井宽度={stair.stair_well_width}mm"
        self.add_paragraph(doc_par)
        doc_par.context = "平面类型：单跑"
        doc_par.is_bold = False
        self.add_paragraph(doc_par)
        self.add_blank_paragraph()

        # 荷载信息
        doc_par = DocParagraph("荷载信息：")
        doc_par.style = self.body_style
        doc_par.first_line_indent = indent
        doc_par.font_size = font_size
        self.add_paragraph(doc_par)

        doc_par = DocParagraph(
            f"附加恒荷载={self.load_param.append_dead_load:.3f}kN/m^{{2}} 活荷载={self.load_param.live_load:.3f}kN/m^{{2}}"
        )
        doc_par.style = self.body_style
        doc_par.first_line_indent = indent * 2
        doc_par.font_size = font_size
        self.add_paragraph(doc_par)
        doc_par.context = f"恒载分项系数：{self.load_param.dead_load_coef}；活载分项系数：{self.load_param.live_load_coef}；活载调整系数：γ_{{L}} ={self.load_param.live_load_adjust_coef:.2f}"
        self.add_paragraph(doc_par)
        doc_par.context = (
            f"活载准永久值系数：{self.load_param.live_load_permenent_coef}"
        )
        self.add_paragraph(doc_par)
        doc_par.context = f"混凝土等级：C{self.concrete_level}，f_{{c}} ={self.concrete_data[self.concrete_level][0]:.2f}MPa"
        self.add_paragraph(doc_par)
        doc_par.context = f"混凝土容重：{self.concrete_density:.2f}kN/mm^{{3}}"
        self.add_paragraph(doc_par)
        doc_par.context = f"配筋调整系数：{self.rebar_area_adjust_coef:.2f}；纵筋保护层厚度：c={self.cover_thickness}mm"
        self.add_paragraph(doc_par)
        doc_par.context = (
            f"梯板纵筋等级：{self.rebar_level}；f_{{y}} ={self.rebar_fy}MPa"
        )
        self.add_paragraph(doc_par)
        doc_par.context = (
            f"梯梁纵筋等级：{self.rebar_level}；f_{{y}} ={self.rebar_fy}MPa"
        )
        self.add_paragraph(doc_par)
        doc_par.context = (
            f"梯梁箍筋等级：{self.rebar_level}；f_{{y}} ={self.rebar_fy}MPa"
        )
        self.add_paragraph(doc_par)

        # 验算信息
        doc_par = DocParagraph("验算信息：")
        doc_par.style = self.body_style
        doc_par.first_line_indent = indent
        doc_par.font_size = font_size
        self.add_paragraph(doc_par)

        doc_par = DocParagraph(
            f"挠度限值：L_{{0}} /{self.displacement_limit}；裂缝限值：{self.crack_width_limit}mm"
        )
        doc_par.style = self.body_style
        doc_par.first_line_indent = indent * 2
        doc_par.font_size = font_size
        self.add_paragraph(doc_par)

        # 计算要求
        doc_par = DocParagraph("计算要求：")
        doc_par.style = self.body_style
        doc_par.first_line_indent = indent
        doc_par.font_size = font_size
        self.add_paragraph(doc_par)

        doc_par = DocParagraph(
            "(1)楼梯板计算；(2)平台梁计算；(3)平台板裂缝验算；(4)平台板挠度验算；"
        )
        doc_par.style = self.body_style
        doc_par.first_line_indent = indent * 2
        doc_par.font_size = font_size
        self.add_paragraph(doc_par)

    def add_2_load_and_calculate(self):
        index = self.current_index
        indent = self.indent
        font_size = self.normal_font_size
        # 添加一个二级标题
        doc_par = DocParagraph(f"{index}.2 荷载与内力计算")
        doc_par.style = self.body_style
        doc_par.first_line_indent = 0
        doc_par.font_size = font_size
        doc_par.is_bold = False
        doc_par.par_level = 2
        self.add_title(doc_par)

        doc_par = DocParagraph("(1)荷载计算：")
        doc_par.style = self.body_style
        doc_par.first_line_indent = 0
        doc_par.font_size = font_size
        self.add_paragraph(doc_par)
        # 标准值
        doc_par.context = "标准值(qk)："
        doc_par.first_line_indent = indent
        self.add_paragraph(doc_par)
        doc_par.context = self.get_load_text(LoadCalulateType.qk)
        doc_par.first_line_indent = indent * 2
        self.add_paragraph(doc_par)
        # 设计值
        doc_par.context = "设计值(q)："
        doc_par.first_line_indent = indent
        self.add_paragraph(doc_par)
        doc_par.context = self.get_load_text(LoadCalulateType.q)
        doc_par.first_line_indent = indent * 2
        self.add_paragraph(doc_par)
        # 准永久值
        doc_par.context = "准永久值(qe)："
        doc_par.first_line_indent = indent
        self.add_paragraph(doc_par)
        doc_par.context = self.get_load_text(LoadCalulateType.qe)
        doc_par.first_line_indent = indent * 2
        self.add_paragraph(doc_par)

        doc_par.context = "(2)内力计算："
        doc_par.first_line_indent = 0
        self.add_paragraph(doc_par)
        # 标准值
        doc_par.context = "a.楼梯板：矩阵位移法求解。"
        doc_par.first_line_indent = indent
        self.add_paragraph(doc_par)

        self.add_blank_paragraph()

    def get_load_text(self, type: LoadCalulateType):
        if type == LoadCalulateType.qk:
            Q1 = self.left_slab_load.qk
            Q2 = self.main_slab_load.qk
            Q3 = self.right_slab_load.qk
        elif type == LoadCalulateType.q:
            Q1 = self.left_slab_load.q
            Q2 = self.main_slab_load.q
            Q3 = self.right_slab_load.q
        elif type == LoadCalulateType.qe:
            Q1 = self.left_slab_load.qe
            Q2 = self.main_slab_load.qe
            Q3 = self.right_slab_load.qe
        else:
            Q1 = Q2 = Q3 = 0

        if self.current_stair.stair_type == "AT":
            return f"斜梯段：{Q2:.2f}kN/m"
        elif self.current_stair.stair_type == "BT":
            return f"左延伸段：{Q1:.2f}kN/m\t斜梯段：{Q2:.2f}kN/m"
        elif self.current_stair.stair_type == "CT":
            return f"斜梯段：{Q2:.2f}kN/m\t右延伸段：{Q3:.2f}kN/m"
        elif self.current_stair.stair_type == "DT":
            return (
                f"左延伸段：{Q1:.2f}kN/m\t斜梯段：{Q2:.2f}kN/m\t右延伸段：{Q3:.2f}kN/m"
            )
        else:
            return "未知踏步类型，暂无数据"

    def add_3_result(self):
        index = self.current_index
        indent = self.indent
        font_size = self.normal_font_size
        # 添加一个二级标题
        doc_par = DocParagraph(f"{index}.3 计算结果")
        doc_par.style = self.body_style
        doc_par.first_line_indent = 0
        doc_par.font_size = font_size
        doc_par.is_bold = False
        doc_par.par_level = 2
        self.add_title(doc_par)

        doc_par = DocParagraph("计算说明：")
        doc_par.style = self.body_style
        doc_par.first_line_indent = 0
        doc_par.font_size = font_size
        self.add_paragraph(doc_par)

        doc_par.context = "(a)简化方法：取板沿着宽度方向单位长度的板带"
        doc_par.first_line_indent = indent
        self.add_paragraph(doc_par)

        doc_par.context = "(b)计算方法：矩阵位移法"
        self.add_paragraph(doc_par)

        doc_par.context = "单位说明："
        doc_par.first_line_indent = 0
        self.add_paragraph(doc_par)

        doc_par.context = "弯        矩：kN·m\t\t剪        力：kN/m\t\t挠        度：mm"
        doc_par.first_line_indent = indent
        self.add_paragraph(doc_par)
        doc_par.context = r"纵筋面积：mm^{2}/m\t\t截面尺寸：mm×mm\t\t裂        缝：mm"
        self.add_paragraph(doc_par)

        self.add_blank_paragraph()
        doc_par.context = "板段配筋计算结果："
        doc_par.first_line_indent = 0
        self.add_paragraph(doc_par)
        doc_par.context = "---------------------------------------------------------------------------------------------"
        self.add_paragraph(doc_par)

        self.insert_calculate_table()

        # self.insert_all_pictures()

        self.add_blank_paragraph()

    def insert_calculate_table(self):
        table_index = 1
        ft = self.concrete_data[self.concrete_level][1]
        if (
            self.current_stair.stair_type == "BT"
            or self.current_stair.stair_type == "DT"
        ):
            left_table = DocTable(8, 4)
            left_table.merge_cells(3, 1, 3, 3)
            left_table.set_table_title(
                f"计算板段-{table_index}(左延伸段)：截面B×H=1000mm×{self.current_stair.left_thick}mm"
            )
            moments = self.current_stair.get_left_slab_table_moments()
            shears = self.current_stair.get_left_slab_table_shears()
            shear_context = self.current_stair.get_shear_validate(
                "left", ft, self.cover_thickness
            )

            table_context = [
                ["截面位置", "左", "中", "右"],
                [
                    "弯矩(M)",
                    f"{moments[0]:.2f}",
                    f"{moments[1]:.2f}",
                    f"{moments[2]:.2f}",
                ],
                ["剪力(V)", f"{shears[0]:.2f}", f"{shears[1]:.2f}", f"{shears[2]:.2f}"],
                ["抗剪截面验算", shear_context],
                ["上部计算纵筋As'", "", "", ""],
                ["下部计算纵筋As", "", "", ""],
                ["上部纵筋实配", "", "", ""],
                ["下部纵筋实配", "", "", ""],
            ]
            left_table.set_table_context(table_context)
            self.add_table(left_table)
            table_index += 1

        cal_table = DocTable(13, 4)
        for i in [3, 8, 9, 11, 12]:
            cal_table.merge_cells(i, 1, i, 3)
        cal_table.set_table_title(
            f"计算板段-{table_index}(斜梯段)：截面B×H=1000mm×{self.current_stair.main_thick}mm"
        )
        moments = self.current_stair.get_main_table_moments()
        shears = self.current_stair.get_main_table_shears()
        shear_context = self.current_stair.get_shear_validate(
            "main", ft, self.cover_thickness
        )

        table_context = [
            ["截面位置", "左", "中", "右"],
            ["弯矩(M)", f"{moments[0]:.2f}", f"{moments[1]:.2f}", f"{moments[2]:.2f}"],
            ["剪力(V)", f"{shears[0]:.2f}", f"{shears[1]:.2f}", f"{shears[2]:.2f}"],
            ["抗剪截面验算", shear_context],
            ["上部计算纵筋As'", "", "", ""],
            ["下部计算纵筋As", "", "", ""],
            ["上部纵筋实配", "", "", ""],
            ["下部纵筋实配", "", "", ""],
            ["挠度限值", f"[f]={self.crack_width_limit:.2f}mm"],
            ["挠度验算结果", "满足"],
            ["裂缝宽度", "", "", ""],
            ["裂缝限值", f"[ω]={self.crack_width_limit:.2f}mm"],
            ["裂缝验算结果", ""],
        ]
        cal_table.set_table_context(table_context)
        self.add_table(cal_table)
        table_index += 1

        if (
            self.current_stair.stair_type == "CT"
            or self.current_stair.stair_type == "DT"
        ):
            right_table = DocTable(8, 4)
            right_table.merge_cells(3, 1, 3, 3)
            right_table.set_table_title(
                f"计算板段-{table_index}(右延伸段)：截面B×H=1000mm×{self.current_stair.right_thick}mm"
            )

            moments = self.current_stair.get_right_slab_table_moments()
            shears = self.current_stair.get_right_slab_table_shears()
            shear_context = self.current_stair.get_shear_validate(
                "right", ft, self.cover_thickness
            )

            table_context = [
                ["截面位置", "左", "中", "右"],
                [
                    "弯矩(M)",
                    f"{moments[0]:.2f}",
                    f"{moments[1]:.2f}",
                    f"{moments[2]:.2f}",
                ],
                ["剪力(V)", f"{shears[0]:.2f}", f"{shears[1]:.2f}", f"{shears[2]:.2f}"],
                ["抗剪截面验算", shear_context],
                ["上部计算纵筋As'", "", "", ""],
                ["下部计算纵筋As", "", "", ""],
                ["上部纵筋实配", "", "", ""],
                ["下部纵筋实配", "", "", ""],
            ]
            right_table.set_table_context(table_context)

            self.add_table(right_table)
            table_index += 1

    def create(self):
        self.add_first_part()
        for i in range(len(self.stair_list)):
            stair = self.stair_list[i]
            self.add_single_stair(stair, i + 1)
            if i < len(self.stair_list) - 1:
                self.add_page_break()

    def save_to_file(self, path):
        self.save(path)
