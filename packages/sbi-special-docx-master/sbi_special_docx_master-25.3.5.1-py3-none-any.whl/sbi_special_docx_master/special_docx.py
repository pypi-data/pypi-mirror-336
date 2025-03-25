import base64
import os
from io import BytesIO
import copy
from typing import Dict, List, Optional, Any, TypeAlias

from docx.document import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from PIL import Image

from .logger import logger
from .validator import validate_data_pydantic
from .styles import DocEditorEmpty

# Set the default image height in inches, configurable via environment variable.
SBIDOCX_PICHEIGHT = float(os.environ.get("SBIDOCX_PICHEIGHT", 2.5))
# Set the block name for the document section, configurable via environment variable.
SBIDOCX_BLOCKNAME = os.getenv("SBIDOCX_BLOCKNAME", "ІНФОРМАЦІЯ ОТ")

InputDataAlias: TypeAlias = Dict[str, List[Dict[str, Any]]]


class AddDocx:
    """
    Class for working with specific blocks in a document.
    It adds textual and image information to a Document object.

    Ensures that the original document is returned if an error occurs during modifications.

    :ivar _orig_doc: DocxDocument - The original document.
    :ivar info: dict - A dictionary containing information to be added.
    :ivar _new_doc: Optional[DocxDocument] - The new document after modifications (if successful).
    """

    def __init__(self, doc: DocxDocument, info: InputDataAlias) -> None:
        """
        Initializes the AddDocx object.

        :param doc: The document object to which information will be added.
        :param info: A dictionary containing the information for addition.
        :raises TypeError: If 'doc' is not an instance of DocxDocument.
        :raises ValueError: If 'info' is not a dictionary.
        """
        if not isinstance(doc, DocxDocument):
            error = f"Received document of type {type(doc)}, expected docx.Document."
            logger.critical(error)
            raise TypeError(error)

        if not isinstance(info, dict):
            logger.critical("No data provided for adding to the document or the data format is incorrect.")
            raise ValueError("Invalid input data format.")

        self._orig_doc: DocxDocument = doc  # Store the original document as a private variable.
        self.info: InputDataAlias = info  # Information provided as a dictionary.
        self._new_doc = None  # To store the new document after modifications.

        # Validate the input data using the pydantic validator.
        validated_items = validate_data_pydantic(self.info)

        try:
            self._new_doc = self._add_info(validated_items)
        except Exception as e:
            logger.error(f"Error {e} while adding information block")

    @property
    def document(self) -> DocxDocument:
        """
        Returns the document instance with the required modifications,
        or the original document if the modified version could not be generated.
        """
        return self._new_doc if self._new_doc is not None else self._orig_doc

    @property
    def has_succeeded(self) -> bool:
        """
        Returns True if the modified document was successfully created,
        or False if the document could not be generated.
        """
        return self._new_doc is not None

    def _add_info(self, info: List):
        """
        Adds textual and image information to the document.

        :param info: A list of dictionaries containing titles, content, and images.
        """
        # If there is no information to add, exit the function.
        if not info:
            return

        # Clone the original document to attempt modifications without affecting the original.
        doc = copy.deepcopy(self._orig_doc)
        doc = DocEditorEmpty(doc).document
        # Add the block header using the configured block name.
        doc.add_paragraph(SBIDOCX_BLOCKNAME, style="left_header")
        for item in info:
            # Add a paragraph for the title with a bullet style.
            title_paragraph = doc.add_paragraph(style='text_base')
            run = title_paragraph.add_run(item['title'])
            run.bold = True  # Make the title bold

            content = item['content']
            # Check if content is a list of paragraphs.
            if isinstance(content, list):
                for entry in content:
                    doc.add_paragraph(text=entry, style='text_base')
            else:
                # If content is not a list, add it directly.
                doc.add_paragraph(content, style='text_base')
            doc.add_paragraph()  # Add an empty paragraph for spacing

            # Attempt to add images from the provided items.
            try:
                for photo in item['images']:
                    # Assumes that each photo dict has a 'file' key containing a Base64 string.
                    doc = self._add_image_from_base64(photo['file'], SBIDOCX_PICHEIGHT, doc)
            except Exception as e:
                logger.info(f"Provided file is not a photo: {e}")
        return doc

    @staticmethod
    def _add_image_from_base64(base64_string: str, height: float, doc: DocxDocument):
        """
        Adds an image from a Base64 string to the document.

        :param base64_string: The image string in Base64 format.
        :param height: The height of the image in inches.
        """
        try:
            # Decode the Base64 string into bytes.
            image_bytes = base64.b64decode(base64_string)
            image_stream = BytesIO(image_bytes)

            # Example code for image processing:
            try:
                image = Image.open(image_stream)
                # Check image integrity.
                image.verify()
                # After calling verify(), reset the file pointer and reopen the image.
                image_stream.seek(0)
                image = Image.open(image_stream)

                # Convert the image into a supported format if needed.
                supported_formats = ["JPEG", "PNG"]
                if image.format not in supported_formats:
                    converted_stream = BytesIO()
                    image.save(converted_stream, format="JPEG")
                    converted_stream.seek(0)
                    image_stream = converted_stream
            except Exception as e:
                logger.warning(f"Error {e} processing image with Pillow")

            # Add the image to the document and center it.
            paragraph = doc.add_paragraph(style='text_base')
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run()
            run.add_picture(image_stream, height=Inches(height))
            doc.add_paragraph()  # Add spacing after the image.
            logger.info("Image for the separate information added successfully.")
        except Exception as e:
            logger.warning(f"Failed to add image: {e}")
            logger.warning(f"Length of Base64 string: {len(base64_string)}")

        return doc

    def save_io(self) -> Optional[BytesIO]:
        """
        Saves the document to an in-memory stream (BytesIO).

        :return: A BytesIO stream with the saved document or None if an error occurs.
        """
        try:
            if self._new_doc:
                saving_doc = self._new_doc
            else:
                saving_doc = self._orig_doc
            output_stream = BytesIO()
            saving_doc.save(output_stream)
            output_stream.seek(0)
            logger.info("Document saving completed successfully.")
            return output_stream
        except Exception as e:
            logger.error(f"Error saving the document: {e}")
            return None

    def save(self, filename: str):
        """
        Saves the document to a file on disk.

        :param filename: The path to the file where the document will be saved.
        """
        try:
            if self._new_doc:
                self._new_doc.save(filename)
                logger.info(f"Document successfully saved as {filename}")
            else:
                self._orig_doc.save(filename)
                logger.info(f"Document successfully saved as {filename}")
        except Exception as e:
            logger.error(f"Error saving the document: {e}")
