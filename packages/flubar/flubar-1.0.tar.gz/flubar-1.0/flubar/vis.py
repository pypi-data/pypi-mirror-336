from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from Bio import SeqIO
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns

def boxplot(input_excel_path, output_image_path, x_label='', y_label='', title='', font_scale=1.2,
                 figsize=(12, 8)):
    """
    Plots a boxplot from an Excel file.

    Parameters:
    - input_excel_path: str, path to the input Excel file.
    - output_image_path: str, path to save the output image.
    - x_label: str, label for the x-axis.
    - y_label: str, label for the y-axis.
    - title: str, title of the plot.
    - font_scale: float, scale factor for the font size.
    - figsize: tuple, figure size in inches.
    """
    # Read the Excel file
    df = pd.read_excel(input_excel_path)

    # Set up the plot
    plt.figure(figsize=figsize)
    sns.set(font_scale=font_scale)

    # Create the boxplot
    boxplot = sns.boxplot(data=df)

    # Customize the plot with labels and title
    boxplot.set_xlabel(x_label, fontsize=14)
    boxplot.set_ylabel(y_label, fontsize=14)
    boxplot.set_title(title, fontsize=16)

    # Save the plot
    plt.savefig(output_image_path, format=output_image_path.split('.')[-1], dpi=300)
    plt.show()

def onedim(input_file, output_file):
    # Determine the file suffix to choose the processing method
    file_suffix = input_file.split('.')[-1].lower()
    seq_records = []

    # Read the file content
    if file_suffix in ['doc', 'docx']:
        doc = Document(input_file)
        seq_records = [{'id': 'Document Text', 'seq': '\n'.join([p.text for p in doc.paragraphs])}]
    elif file_suffix == 'txt':
        with open(input_file, 'r') as file:
            seq_records = [{'id': 'Text File', 'seq': file.read()}]
    elif file_suffix == 'fasta':
        seq_records = list(SeqIO.parse(input_file, 'fasta'))
    else:
        raise ValueError("Unsupported file format")

    # Replace bases and record the replacement positions and base types
    base_replacements = {
        'A': ('|', 'Showcard Gothic', 14, RGBColor(0x80, 0x76, 0xA3)),  # Size 14 font roughly corresponds to 12pt
        'T': ('|', 'Showcard Gothic', 14, RGBColor(0xFF, 0x00, 0x00)),
        'G': ('|', 'Showcard Gothic', 16, RGBColor(0x56, 0x98, 0xC3)),  # Size 16 font roughly corresponds to 16pt
        'C': ('|', 'Showcard Gothic', 16, RGBColor(0x00, 0xB0, 0x50)),
    }

    # Create a new Word document
    doc = Document()
    # Set the overall style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Showcard Gothic'

    # Add sequences and their names to the document
    for rec in seq_records:
        # Add sequence ID
        doc.add_paragraph(rec.id, style='Heading 1')

        # Add sequence content
        seq = str(rec.seq)
        for line in seq.split('\n'):
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            for char in line:
                run = p.add_run(char)
                if char in base_replacements:
                    replacement, font_name, font_size, color = base_replacements[char]
                    run.text = replacement
                    run.font.name = font_name
                    run.font.size = Pt(font_size)
                    run.font.color.rgb = color

    # Save the new Word document
    doc.save(output_file)

def heatmap(input_excel_path, output_image_path, cmap='coolwarm', annot=True, fmt=".3f", font_scale=1.2,
                 figsize=(10, 8)):
    """
    Plots a heatmap from an Excel file.

    Parameters:
    - input_excel_path: str, path to the input Excel file.
    - output_image_path: str, path to save the output image.
    - cmap: str, colormap for the heatmap.
    - annot: bool, whether to annotate the heatmap.
    - fmt: str, string formatting code for annotations.
    - font_scale: float, scale factor for the font size.
    - figsize: tuple, figure size in inches.
    """
    # Read the Excel file
    df = pd.read_excel(input_excel_path, index_col=0)

    # Set up the plot
    plt.figure(figsize=figsize)
    sns.set(font_scale=font_scale)
    heatmap = sns.heatmap(df, cmap=cmap, annot=annot, fmt=fmt, linewidths=.5, cbar_kws={'label': 'Values'})

    # Add title and adjust layout
    heatmap.set_title('Heatmap of Excel Data', fontsize=16)
    plt.xticks(rotation=45, ha='right')
    plt.yticks(rotation=0)
    plt.tight_layout()

    # Save the plot
    plt.savefig(output_image_path, format=output_image_path.split('.')[-1], dpi=300)
    plt.show()


# # 1维码可视化
# input_file = 'barcodes_input.fasta'  # Replace with the actual file path
# output_file = 'output.docx'  # Replace with the desired output file path
# one_dim_code(input_file, output_file)
#
# # 热图绘制
# input_excel_path = 'heatmap_input.xlsx'  # Replace with the actual input Excel file path
# output_image_path = 'heatmap.svg'  # Replace with the desired output image file path, can be 'heatmap.svg', 'heatmap.png', etc.
# plot_heatmap(input_excel_path, output_image_path, cmap='viridis', annot=True, fmt=".2f", font_scale=1.5, figsize=(12, 10))
#
# # 箱线图绘制
# input_excel_path = 'boxplot_input.xlsx'  # Replace with the actual input Excel file path
# output_image_path = 'boxplot.svg'  # Replace with the desired output image file path, can be 'boxplot.svg', 'boxplot.png', etc.
# plot_boxplot(input_excel_path, output_image_path, x_label='Samples', y_label='Values', title='Boxplot of Excel Data', font_scale=1.5, figsize=(14, 10))
