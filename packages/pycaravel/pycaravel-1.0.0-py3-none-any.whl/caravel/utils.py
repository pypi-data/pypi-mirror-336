###############################################################################
# NSAp - Copyright (C) CEA, 2019
# Distributed under the terms of the CeCILL-B license, as published by
# the CEA-CNRS-INRIA. Refer to the LICENSE file or to
# http://www.cecill.info/licences/Licence_CeCILL-B_V1-en.html
# for details.
###############################################################################


"""
This module provides common utilities.
"""

# Imports
import os
import re
import json
import time
import functools
from docx import Document


def export_report(report, timestamp, outfile):
    """ Export the report in docx format.

    Parameters
    ----------
    report: str
        a dictionary with recursive keys and values.
    timestamp: str
        a timestamp.
    outfile: str
        the path to the generated docx file.
    """
    document = Document()
    section = document.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.text = f"NeuroSpin\tReporting\t{timestamp}"
    paragraph = document.add_paragraph(
        f"\n\n\n\nYou will find below the report generated on '{timestamp}'. "
        "If you have any questions please use the contact mail: rlink@cea.fr.")
    for family, family_item in report.items():
        document.add_heading(family.replace(".", " ").title())
        for validator, validator_item in family_item.items():
            split_validator = re.findall("[A-Z][^A-Z]*", validator)
            document.add_heading(" ".join(split_validator))
            paragraph = document.add_paragraph(
                "\n\n Below the table summarizing the errors.\n\n")
            for key, values in validator_item.items():
                table = document.add_table(rows=len(values), cols=2)
                cell = table.cell(0, 0)
                cell.text = key
                for idx, val in enumerate(values):
                    cell = table.cell(idx, 1)
                    cell.text = val
    document.save(outfile)


def monitor(func):
    """ A decorator to monitor function and log its status in a root directory.
    The input function parameters must be set via the 'CARAVEL_ROOT'
    and 'CARAVEL_NAME' environement variables.
    """
    root = os.environ.get("CARAVEL_ROOT", None)
    name = os.environ.get("CARAVEL_NAME", None)
    is_monitor = root is not None and name is not None
    if is_monitor:
        assert os.path.isdir(root), root

    @functools.wraps(func)
    def decorated(*args, **kwargs):
        try:
            tic = time.time()
            res = func(*args, **kwargs)
            toc = time.time()
            info = {
                "status": "healthy",
                "duration": toc - tic
            }
            if is_monitor:
                with open(os.path.join(root, f"{name}.json"), "wt") as of:
                    json.dump(info, of, indent=4)
            return res
        except Exception as e:
            info = {
                "status": "dead"
            }
            if is_monitor:
                with open(os.path.join(root, f"{name}.json"), "wt") as of:
                    json.dump(info, of, indent=4)
            raise e
    return decorated
